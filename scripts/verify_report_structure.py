from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def find_target() -> Path:
    base = Path.home() / "Downloads" / "Telegram Desktop"
    candidates = [
        p
        for p in base.glob("*_tables_conclusions.docx")
        if not p.name.startswith("~$")
    ]
    if not candidates:
        raise FileNotFoundError("No final tables/conclusions DOCX found.")
    return max(candidates, key=lambda p: p.stat().st_mtime)


def paragraph_text(element) -> str:
    if element.tag != qn("w:p"):
        return ""
    return "".join(t.text or "" for t in element.iter(qn("w:t"))).strip()


def paragraph_style_id(element) -> str:
    if element.tag != qn("w:p"):
        return ""
    p_pr = element.find(qn("w:pPr"))
    if p_pr is None:
        return ""
    p_style = p_pr.find(qn("w:pStyle"))
    if p_style is None:
        return ""
    return p_style.get(qn("w:val"), "")


def is_heading(element) -> bool:
    text = paragraph_text(element)
    style_id = paragraph_style_id(element).lower()
    if style_id.startswith("heading") or style_id.startswith("заголовок"):
        return True
    if text in {"ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ", "СПИСОК ИСПОЛЬЗУЕМЫХ ИСТОЧНИКОВ"}:
        return True
    if text.startswith(("1 ", "2 ", "3 ", "1.", "2.", "3.")):
        return True
    if text.startswith("Выводы по разделу"):
        return True
    return False


def conclusion_stats(document: Document) -> dict[str, tuple[int, int]]:
    elements = list(document._body._element)
    stats: dict[str, tuple[int, int]] = {}
    for section_number in ("1", "2", "3"):
        start = None
        for idx, element in enumerate(elements):
            if paragraph_text(element).startswith(f"Выводы по разделу {section_number}"):
                start = idx
                break
        if start is None:
            stats[section_number] = (0, 0)
            continue
        paragraphs = 0
        words = 0
        for element in elements[start + 1 :]:
            if is_heading(element):
                break
            text = paragraph_text(element)
            if text:
                paragraphs += 1
                words += len(text.split())
        stats[section_number] = (paragraphs, words)
    return stats


def table_audit(document: Document) -> dict[str, int]:
    stats = {
        "tables": len(document.tables),
        "non_left_paragraphs": 0,
        "non_bw_shading": 0,
        "non_black_borders": 0,
    }
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = tc_pr.find(qn("w:shd"))
                if shd is not None:
                    fill = (shd.get(qn("w:fill")) or "").upper()
                    if fill not in {"", "AUTO", "FFFFFF"}:
                        stats["non_bw_shading"] += 1
                borders = tc_pr.find(qn("w:tcBorders"))
                if borders is not None:
                    for border in borders:
                        color = (border.get(qn("w:color")) or "").upper()
                        if color not in {"", "AUTO", "000000"}:
                            stats["non_black_borders"] += 1
                for paragraph in cell.paragraphs:
                    jc = paragraph._p.pPr.jc if paragraph._p.pPr is not None else None
                    val = jc.val if jc is not None else None
                    normalized = str(val).lower() if val is not None else ""
                    if val is not None and "left" not in normalized and "start" not in normalized:
                        stats["non_left_paragraphs"] += 1
    return stats


def endings_audit(document: Document) -> list[str]:
    elements = list(document._body._element)
    issues: list[str] = []
    last_heading = None
    last_content = None
    for element in elements:
        if element.tag == qn("w:sectPr"):
            continue
        if is_heading(element):
            if last_heading and last_content in {"table", "figure"}:
                issues.append(f"{last_heading} ends with {last_content}")
            last_heading = paragraph_text(element)
            last_content = None
            continue
        if element.tag == qn("w:tbl"):
            last_content = "table"
            continue
        text = paragraph_text(element)
        if text:
            if text.lower().startswith("рисунок"):
                last_content = "figure"
            else:
                last_content = "text"
    if last_heading and last_content in {"table", "figure"}:
        issues.append(f"{last_heading} ends with {last_content}")
    return issues


def main() -> None:
    target = find_target()
    document = Document(target)
    print(f"TARGET={target}")
    print(f"CONCLUSIONS={conclusion_stats(document)}")
    print(f"TABLE_AUDIT={table_audit(document)}")
    print(f"ENDING_ISSUES={endings_audit(document)}")


if __name__ == "__main__":
    main()
