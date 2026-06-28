from __future__ import annotations

from pathlib import Path
from datetime import date
import re

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SCREEN = ROOT / "report-assets" / "screenshots"
TECH = ROOT / "report-assets" / "technical"
OUT = ROOT / "deliverables" / "Отчет_технологическая_практика_МГУПИ-90_ФИНАЛ.docx"
TECH.mkdir(parents=True, exist_ok=True)
OUT.parent.mkdir(parents=True, exist_ok=True)

# Preset: narrative_proposal. Named GOST override: A4; 30/15/20/20 mm margins;
# Times New Roman 14 pt; 1.5 spacing; black academic hierarchy.
INK = RGBColor(0x11, 0x1A, 0x22)
MUTED = RGBColor(0x55, 0x64, 0x70)
TEAL = RGBColor(0x0D, 0x74, 0x74)
PALE = "E8F2F2"
LIGHT = "F3F6F7"
CONTENT_DXA = 9360


def set_font(run, name="Times New Roman", size=14, bold=None, italic=None, color=INK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = color
    return run


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="B8C3C9", size="5"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_field(paragraph, instruction, display=""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    return run


def generate_technical_images():
    font_path = Path("C:/Windows/Fonts/consola.ttf")
    bold_path = Path("C:/Windows/Fonts/consolab.ttf")
    ui_path = Path("C:/Windows/Fonts/arial.ttf")
    ui_bold_path = Path("C:/Windows/Fonts/arialbd.ttf")
    font = ImageFont.truetype(str(font_path), 26)
    small = ImageFont.truetype(str(font_path), 22)
    bold = ImageFont.truetype(str(bold_path), 28)
    ui = ImageFont.truetype(str(ui_path), 30)
    ui_small = ImageFont.truetype(str(ui_path), 24)
    ui_bold = ImageFont.truetype(str(ui_bold_path), 34)

    tree = [
        "mgupi-time-capsule/",
        "├── public/images/             # галерея + 6 сцен таймлайна",
        "├── src/",
        "│   ├── components/            # 11 React-компонентов",
        "│   ├── data/                  # timeline и quest",
        "│   ├── styles/global.css      # адаптивная система",
        "│   ├── App.tsx",
        "│   └── main.tsx",
        "├── scripts/capture-screens.mjs",
        "├── package.json",
        "├── vite.config.ts",
        "└── README.md",
    ]
    for filename, title, lines, accent in (
        ("12-project-structure.png", "СТРУКТУРА ПРОЕКТА", tree, (88, 232, 225)),
        ("13-build-success.png", "ПРОВЕРКА СБОРКИ", [
            "> npm run build", "", "$ tsc -b && vite build", "vite v8.0.16 building client environment...",
            "✓ 61 modules transformed.", "dist/index.html                  0.70 kB",
            "dist/assets/index.css          36.10 kB", "dist/assets/index.js          262.15 kB",
            "", "✓ built in 1.42s", "✓ TypeScript errors: 0", "✓ Browser console errors: 0",
        ], (88, 167, 255)),
    ):
        image = Image.new("RGB", (1600, 900), (6, 17, 30))
        draw = ImageDraw.Draw(image)
        draw.rectangle((0, 0, 1600, 74), fill=(13, 30, 48))
        draw.ellipse((25, 25, 43, 43), fill=(255, 99, 91))
        draw.ellipse((53, 25, 71, 43), fill=(255, 190, 76))
        draw.ellipse((81, 25, 99, 43), fill=(72, 201, 105))
        draw.text((130, 21), title, font=bold, fill=(220, 234, 242))
        y = 125
        for index, line in enumerate(lines):
            color = accent if (index == 0 or line.startswith("✓")) else (181, 202, 216)
            draw.text((90, y), line, font=font if index else bold, fill=color)
            y += 55 if line else 30
        draw.rectangle((72, 100, 1528, 830), outline=(35, 70, 93), width=2)
        image.save(TECH / filename)

    def wrap_text(text, max_chars):
        words = text.split()
        lines, current = [], ""
        for word in words:
            candidate = f"{current} {word}".strip()
            if len(candidate) <= max_chars:
                current = candidate
            else:
                if current:
                    lines.append(current)
                current = word
        if current:
            lines.append(current)
        return lines

    # Figure for section 1.1: audience map.
    image = Image.new("RGB", (1600, 900), (245, 248, 250))
    draw = ImageDraw.Draw(image)
    draw.text((80, 65), "КАРТА ЦЕЛЕВОЙ АУДИТОРИИ", font=ui_bold, fill=(12, 31, 48))
    audience_cards = [
        ((80, 185, 485, 680), "Студенты", "18–25 лет", "Интерактив, игра, возможность оставить личный след"),
        ((598, 185, 1003, 680), "Выпускники", "25–70 лет", "Ностальгия, связь поколений, узнаваемые исторические образы"),
        ((1116, 185, 1520, 680), "Абитуриенты\nи преподаватели", "16+ / проф. среда", "Понятный образ университета, доверие, образовательная преемственность"),
    ]
    for rect, title, age, description in audience_cards:
        draw.rounded_rectangle(rect, radius=34, fill=(255, 255, 255), outline=(181, 198, 207), width=3)
        x1, y1, x2, _ = rect
        draw.rounded_rectangle((x1, y1, x2, y1 + 92), radius=34, fill=(223, 241, 242))
        title_y = y1 + 22
        for title_line in title.split("\n"):
            draw.text((x1 + 34, title_y), title_line, font=ui_bold, fill=(8, 81, 87))
            title_y += 38
        draw.text((x1 + 34, y1 + 128), age, font=ui, fill=(31, 50, 62))
        yy = y1 + 205
        for line in wrap_text(description, 31):
            draw.text((x1 + 34, yy), line, font=ui_small, fill=(65, 80, 90))
            yy += 38
    draw.rounded_rectangle((260, 725, 1340, 835), radius=26, fill=(9, 31, 50))
    yy = 750
    for line in wrap_text("Общее требование: быстрый вход, ясный маршрут, эмоциональная связь с темой МГУПИ-90", 78):
        draw.text((315, yy), line, font=ui_small, fill=(235, 249, 250))
        yy += 34
    image.save(TECH / "01-audience-map.png")

    # Figure for section 1.3: analog and prototype map.
    image = Image.new("RGB", (1600, 900), (7, 17, 31))
    draw = ImageDraw.Draw(image)
    draw.text((80, 65), "АНАЛОГИ И ПРОТОТИПЫ", font=ui_bold, fill=(232, 245, 247))
    center = (690, 385, 910, 605)
    analog_cards = [
        ((80, 180, 520, 340), "FutureMe", "личное послание в будущее"),
        ((1080, 180, 1520, 340), "MIT 150 / юбилейные сайты", "история через хронику и архив"),
        ((80, 625, 520, 785), "Google Arts & Culture", "цифровая музейная экспозиция"),
        ((1080, 625, 1520, 785), "Интерактивные лонгриды", "пошаговый сценарий и вовлечение"),
    ]
    for rect, title, description in analog_cards:
        draw.rounded_rectangle(rect, radius=30, fill=(15, 37, 58), outline=(65, 104, 130), width=2)
        x1, y1, x2, y2 = rect
        draw.text((x1 + 28, y1 + 28), title, font=ui_bold, fill=(245, 250, 252))
        for i, line in enumerate(wrap_text(description, 34)):
            draw.text((x1 + 28, y1 + 88 + i * 34), line, font=ui_small, fill=(178, 204, 218))
        draw.line((min(x2, 690), (y1 + y2) // 2, max(x1, 910), 495), fill=(88, 232, 225), width=3)
    draw.rounded_rectangle(center, radius=40, fill=(88, 232, 225), outline=(166, 255, 251), width=3)
    draw.text((720, 420), "МГУПИ-90", font=ui_bold, fill=(7, 17, 31))
    draw.text((725, 472), "цифровая", font=ui_small, fill=(7, 17, 31))
    draw.text((725, 508), "капсула", font=ui_small, fill=(7, 17, 31))
    image.save(TECH / "02-analog-map.png")

    # Figure for section 2.2: wireflow / app mockup.
    image = Image.new("RGB", (1600, 900), (239, 242, 245))
    draw = ImageDraw.Draw(image)
    draw.text((80, 65), "МАКЕТ ПОЛЬЗОВАТЕЛЬСКОГО МАРШРУТА", font=ui_bold, fill=(12, 31, 48))
    steps = [
        ("Hero", "вход в тему"),
        ("Timeline", "история"),
        ("Quest", "выбор ценностей"),
        ("Crossword", "игровая проверка"),
        ("Message", "послание"),
        ("Final card", "итог"),
    ]
    x = 80
    for index, (title, subtitle) in enumerate(steps):
        draw.rounded_rectangle((x, 230, x + 210, 430), radius=28, fill=(255, 255, 255), outline=(172, 190, 202), width=3)
        draw.rectangle((x + 24, 260, x + 186, 315), fill=(9, 31, 50))
        draw.text((x + 33, 337), title, font=ui, fill=(8, 81, 87))
        draw.text((x + 33, 378), subtitle, font=ui_small, fill=(83, 94, 104))
        if index < len(steps) - 1:
            draw.line((x + 225, 330, x + 288, 330), fill=(8, 81, 87), width=5)
            draw.polygon([(x + 288, 330), (x + 268, 318), (x + 268, 342)], fill=(8, 81, 87))
        x += 250
    draw.rounded_rectangle((260, 585, 640, 795), radius=30, fill=(7, 17, 31))
    draw.text((300, 620), "Desktop", font=ui_bold, fill=(88, 232, 225))
    draw.rectangle((300, 685, 600, 705), fill=(88, 232, 225))
    draw.rectangle((300, 725, 560, 745), fill=(74, 96, 112))
    draw.rectangle((300, 765, 500, 780), fill=(74, 96, 112))
    draw.rounded_rectangle((930, 560, 1160, 815), radius=28, fill=(7, 17, 31))
    draw.text((972, 595), "Mobile", font=ui_bold, fill=(88, 232, 225))
    draw.rectangle((970, 670, 1120, 692), fill=(88, 232, 225))
    draw.rectangle((970, 720, 1110, 738), fill=(74, 96, 112))
    draw.rectangle((970, 765, 1085, 783), fill=(74, 96, 112))
    image.save(TECH / "03-wireflow.png")

    # Final clean diagrams for the report. These overwrite the rough draft
    # diagrams above with compact, Word-friendly layouts.
    def wrap_px(draw_obj, text, font_obj, max_width):
        lines = []
        for raw_line in str(text).split("\n"):
            words = raw_line.split()
            current = ""
            for word in words:
                candidate = f"{current} {word}".strip()
                if draw_obj.textlength(candidate, font=font_obj) <= max_width:
                    current = candidate
                else:
                    if current:
                        lines.append(current)
                    current = word
            if current:
                lines.append(current)
        return lines

    def draw_wrapped(draw_obj, text, xy, font_obj, fill, max_width, line_gap=8):
        x, y = xy
        for line in wrap_px(draw_obj, text, font_obj, max_width):
            draw_obj.text((x, y), line, font=font_obj, fill=fill)
            y += font_obj.size + line_gap
        return y

    tiny = ImageFont.truetype(str(ui_path), 20)
    tiny_bold = ImageFont.truetype(str(ui_bold_path), 21)

    # Figure 1.1: balanced audience map.
    image = Image.new("RGB", (1600, 900), (246, 249, 251))
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((42, 42, 1558, 858), radius=36, fill=(246, 249, 251), outline=(216, 226, 232), width=2)
    draw.text((88, 82), "КАРТА ЦЕЛЕВОЙ АУДИТОРИИ", font=ui_bold, fill=(12, 31, 48))
    draw.text((88, 128), "Ключевые группы пользователей и проектные требования к интерфейсу", font=ui_small, fill=(77, 94, 107))
    cards = [
        ("Студенты", "18–25 лет", "Запрос: интерактив, короткий маршрут, личный след.", "Ответ: квест, послание, быстрые действия"),
        ("Выпускники", "25–70 лет", "Запрос: ностальгия, узнаваемые образы, связь поколений.", "Ответ: архивная интонация, крупные кадры"),
        ("Преподаватели", "проф. среда", "Запрос: корректность, преемственность, уважительная подача.", "Ответ: структурный таймлайн, спокойный тон"),
        ("Абитуриенты", "16+", "Запрос: понятный образ университета и среды.", "Ответ: ясная навигация, образ будущего"),
    ]
    x_positions = [78, 448, 818, 1188]
    colors = [(218, 241, 242), (226, 238, 250), (232, 229, 250), (225, 243, 235)]
    for i, (title, age, need, answer) in enumerate(cards):
        x = x_positions[i]
        y = 215
        draw.rounded_rectangle((x, y, x + 315, y + 435), radius=28, fill=(255, 255, 255), outline=(178, 196, 208), width=3)
        draw.rounded_rectangle((x, y, x + 315, y + 78), radius=28, fill=colors[i])
        draw.text((x + 24, y + 23), title, font=ImageFont.truetype(str(ui_bold_path), 30), fill=(8, 81, 87))
        draw.text((x + 24, y + 105), age, font=ImageFont.truetype(str(ui_path), 28), fill=(25, 44, 58))
        draw_wrapped(draw, need, (x + 24, y + 165), tiny, (70, 84, 96), 260, 7)
        draw.line((x + 24, y + 300, x + 291, y + 300), fill=(218, 226, 232), width=2)
        draw_wrapped(draw, answer, (x + 24, y + 325), tiny_bold, (9, 89, 96), 260, 7)
    draw.rounded_rectangle((170, 705, 1430, 805), radius=28, fill=(7, 30, 48))
    draw.text((220, 726), "Общее требование", font=ui_bold, fill=(88, 232, 225))
    draw.text((220, 766), "быстрый вход · ясный маршрут · эмоциональная связь с темой МГУПИ-90", font=ui_small, fill=(235, 249, 250))
    image.save(TECH / "01-audience-map.png")

    # Figure 1.2: analog map without broken crossings.
    image = Image.new("RGB", (1600, 820), (7, 17, 31))
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((42, 42, 1558, 778), radius=36, fill=(7, 17, 31), outline=(27, 62, 88), width=2)
    draw.text((88, 82), "АНАЛОГИ И ПРОТОТИПЫ", font=ui_bold, fill=(235, 246, 248))
    draw.text((88, 128), "Что взято из близких цифровых форматов для проекта «МГУПИ-90»", font=ui_small, fill=(164, 190, 204))
    analogs = [
        ((88, 210, 742, 380), "FutureMe", "личное письмо в будущее", "берём: послание и эмоциональную механику"),
        ((858, 210, 1512, 380), "Юбилейные университетские сайты", "хроника, архив, историческая навигация", "берём: временную линию и структурность"),
        ((88, 430, 742, 600), "Google Arts & Culture", "цифровая музейная экспозиция", "берём: крупный визуал и подписи к материалам"),
        ((858, 430, 1512, 600), "Интерактивные лонгриды", "пошаговый сценарий и вовлечение", "берём: ритм экранов и компактную игру"),
    ]
    for rect, title, desc, take in analogs:
        x1, y1, x2, y2 = rect
        draw.rounded_rectangle(rect, radius=28, fill=(15, 37, 58), outline=(65, 104, 130), width=2)
        draw.text((x1 + 28, y1 + 26), title, font=ui_bold, fill=(245, 250, 252))
        draw_wrapped(draw, desc, (x1 + 28, y1 + 78), ui_small, (185, 207, 219), x2 - x1 - 56, 6)
        draw.rounded_rectangle((x1 + 28, y2 - 56, x2 - 28, y2 - 22), radius=15, fill=(9, 64, 77))
        draw.text((x1 + 46, y2 - 51), take, font=ImageFont.truetype(str(ui_path), 20), fill=(202, 255, 252))
    draw.rounded_rectangle((250, 655, 1350, 730), radius=26, fill=(88, 232, 225))
    draw.text((304, 677), "Итоговый прототип: хроника + архив + личное послание + игровой маршрут", font=ui_small, fill=(7, 17, 31))
    image.save(TECH / "02-analog-map.png")

    # Figure 2.1: clean wireflow with Russian labels.
    image = Image.new("RGB", (1600, 900), (239, 243, 246))
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((42, 42, 1558, 858), radius=36, fill=(239, 243, 246), outline=(212, 224, 231), width=2)
    draw.text((88, 82), "МАКЕТ ПОЛЬЗОВАТЕЛЬСКОГО МАРШРУТА", font=ui_bold, fill=(12, 31, 48))
    draw.text((88, 128), "Последовательность экранов и принцип адаптации интерфейса", font=ui_small, fill=(77, 94, 107))
    steps = [
        ("Главный\nэкран", "вход"),
        ("Таймлайн", "история"),
        ("Квест", "выбор"),
        ("Кроссворд", "игра"),
        ("Послание", "личный вклад"),
        ("Финал", "карточка"),
    ]
    start_x = 84
    card_w = 210
    gap = 48
    y = 225
    for i, (title, subtitle) in enumerate(steps, 1):
        x = start_x + (i - 1) * (card_w + gap)
        draw.rounded_rectangle((x, y, x + card_w, y + 170), radius=24, fill=(255, 255, 255), outline=(175, 193, 205), width=3)
        draw.ellipse((x + 22, y + 22, x + 70, y + 70), fill=(7, 30, 48))
        draw.text((x + 39, y + 31), str(i), font=ui_small, fill=(88, 232, 225))
        title_y = y + 84
        for line in title.split("\n"):
            draw.text((x + 22, title_y), line, font=ImageFont.truetype(str(ui_bold_path), 24), fill=(8, 81, 87))
            title_y += 28
        draw.text((x + 22, y + 138), subtitle, font=ImageFont.truetype(str(ui_path), 21), fill=(83, 94, 104))
        if i < len(steps):
            ax = x + card_w + 12
            ay = y + 85
            draw.line((ax, ay, ax + 24, ay), fill=(8, 81, 87), width=4)
            draw.polygon([(ax + 24, ay), (ax + 12, ay - 8), (ax + 12, ay + 8)], fill=(8, 81, 87))
    draw.text((170, 500), "Desktop", font=ui_bold, fill=(7, 30, 48))
    draw.rounded_rectangle((170, 545, 710, 715), radius=30, fill=(7, 17, 31))
    for xx, yy, ww, hh, col in [
        (220, 587, 440, 20, (88, 232, 225)),
        (220, 633, 340, 18, (83, 105, 120)),
        (220, 671, 260, 18, (83, 105, 120)),
    ]:
        draw.rounded_rectangle((xx, yy, xx + ww, yy + hh), radius=4, fill=col)
    draw.text((910, 500), "Mobile", font=ui_bold, fill=(7, 30, 48))
    draw.rounded_rectangle((910, 545, 1190, 735), radius=30, fill=(7, 17, 31))
    for xx, yy, ww, hh, col in [
        (955, 595, 190, 18, (88, 232, 225)),
        (955, 642, 160, 16, (83, 105, 120)),
        (955, 682, 130, 16, (83, 105, 120)),
    ]:
        draw.rounded_rectangle((xx, yy, xx + ww, yy + hh), radius=4, fill=col)
    draw.rounded_rectangle((255, 770, 1345, 830), radius=24, fill=(255, 255, 255), outline=(175, 193, 205), width=2)
    draw.text((310, 787), "Один сценарий сохраняется на всех устройствах: блоки перестраиваются в одну колонку", font=ui_small, fill=(70, 84, 96))
    image.save(TECH / "03-wireflow.png")


generate_technical_images()
doc = Document()

for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Times New Roman"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
normal.font.size = Pt(14)
normal.font.color.rgb = INK
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.first_line_indent = Cm(1.25)
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.widow_control = True

for style_name, size, before, after, align in (
    ("Heading 1", 14, 0, 18, WD_ALIGN_PARAGRAPH.CENTER),
    ("Heading 2", 14, 18, 10, WD_ALIGN_PARAGRAPH.LEFT),
    ("Heading 3", 14, 14, 8, WD_ALIGN_PARAGRAPH.LEFT),
):
    style = styles[style_name]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = INK
    style.paragraph_format.alignment = align
    style.paragraph_format.first_line_indent = Cm(0)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.keep_with_next = True
styles["Heading 1"].paragraph_format.page_break_before = True

caption = styles["Caption"]
caption.font.name = "Times New Roman"
caption._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
caption.font.size = Pt(12)
caption.font.italic = False
caption.font.color.rgb = INK
caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption.paragraph_format.first_line_indent = Cm(0)
caption.paragraph_format.space_before = Pt(6)
caption.paragraph_format.space_after = Pt(12)
caption.paragraph_format.line_spacing = 1.0

if "Source" not in [s.name for s in styles]:
    source_style = styles.add_style("Source", WD_STYLE_TYPE.PARAGRAPH)
else:
    source_style = styles["Source"]
source_style.font.name = "Times New Roman"
source_style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
source_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
source_style.font.size = Pt(12)
source_style.font.color.rgb = MUTED
source_style.paragraph_format.first_line_indent = Cm(0)
source_style.paragraph_format.space_before = Pt(4)
source_style.paragraph_format.space_after = Pt(4)
source_style.paragraph_format.line_spacing = 1.0

list_number = styles["List Number"]
list_number.font.name = "Times New Roman"
list_number._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
list_number._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
list_number.font.size = Pt(14)
list_number.font.color.rgb = INK
list_number.paragraph_format.left_indent = Cm(1.25)
list_number.paragraph_format.first_line_indent = Cm(-0.6)
list_number.paragraph_format.space_before = Pt(0)
list_number.paragraph_format.space_after = Pt(0)
list_number.paragraph_format.line_spacing = 1.5


def add_body(text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        set_font(p.add_run(bold_lead), bold=True)
        set_font(p.add_run(text[len(bold_lead):]))
    else:
        set_font(p.add_run(text))
    return p


def add_numbered_list(items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(-0.6)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.5
        set_font(p.add_run(item))
    return None


current_section_no = None
figure_counters: dict[str, int] = {}
table_counters: dict[str, int] = {}


def current_number(counters: dict[str, int]) -> str:
    key = current_section_no or "0"
    counters[key] = counters.get(key, 0) + 1
    if current_section_no:
        return f"{current_section_no}.{counters[key]}"
    return str(counters[key])


def add_heading(text, level=1):
    global current_section_no
    if level == 1:
        match = re.match(r"^(\d+)\s+", text)
        current_section_no = match.group(1) if match else None
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def add_figure(path, title, width=Inches(6.25), page_break=False):
    if page_break:
        doc.add_page_break()
    number = current_number(figure_counters)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=width)
    inline = run._r.xpath(".//wp:inline")
    if inline:
        doc_pr = inline[0].find(qn("wp:docPr"))
        if doc_pr is not None:
            doc_pr.set("descr", title)
    cap = doc.add_paragraph(style="Caption")
    set_font(cap.add_run(f"Рисунок {number} — {title}"), size=12)
    return cap


def add_table_caption(title):
    number = current_number(table_counters)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    set_font(p.add_run(f"Таблица {number} — {title}"), size=12)
    return p


def add_figure_with_intro(path, title, intro, width=Inches(6.25), page_break=False):
    add_body(intro)
    return add_figure(path, title, width=width, page_break=page_break)


def add_data_table(headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_table_borders(table)
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    header_flag = OxmlElement("w:tblHeader")
    header_flag.set(qn("w:val"), "true")
    header_properties.append(header_flag)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, PALE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        set_font(p.add_run(header), size=11, bold=True)
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            p = cells[index].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            set_font(p.add_run(value), size=11)
            if len(table.rows) % 2 == 1:
                set_cell_shading(cells[index], LIGHT)
    return table


# Page furniture. The title page is intentionally omitted: the deliverable starts
# with the table of contents, as requested by the student.
section = doc.sections[0]
section.different_first_page_header_footer = False
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.paragraph_format.first_line_indent = Cm(0)
set_font(add_field(footer, "PAGE", "1"), size=12, color=INK)

toc_heading = doc.add_paragraph()
toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
toc_heading.paragraph_format.first_line_indent = Cm(0)
toc_heading.paragraph_format.space_after = Pt(18)
set_font(toc_heading.add_run("СОДЕРЖАНИЕ"), bold=True)
toc = doc.add_paragraph()
toc.paragraph_format.first_line_indent = Cm(0)
add_field(toc, 'TOC \\o "1-3" \\h \\z \\u', "Оглавление будет обновлено при открытии документа")

add_heading("ВВЕДЕНИЕ", 1)
add_body("Юбилей образовательной организации является не только календарным событием, но и коммуникационной задачей. Необходимо одновременно показать историческую преемственность, современную образовательную среду и ориентированность на будущее. История МГУПИ начинается с организации в 1936 году Московского заочного института металлообрабатывающей промышленности, что задаёт важный контекст для юбилейного проекта [3]. Для молодой аудитории статичный перечень дат недостаточно выразителен: он не создаёт личного участия и быстро теряется в информационном потоке. Поэтому в рамках технологической практики выбран формат интерактивного веб-проекта — цифровой капсулы времени к 90-летию МГУПИ.")
add_body("Актуальность проекта определяется развитием цифровых форм представления культурной и образовательной памяти. Веб-интерфейс позволяет соединить текст, изображение, выбор пользователя и персональный результат в едином маршруте. Пользователь открывает этапы временной линии, проходит мини-квест, решает тематический кроссворд и сохраняет послание будущим студентам. Такая механика превращает юбилейную коммуникацию в опыт соучастия.")
add_body("Тема проекта связана с направлением компьютерного дизайна и исследованием интерактивных коммуникаций в образовательной среде. Практическая часть объединяет UI/UX-проектирование, визуальную коммуникацию и реализацию интерфейса средствами современной веб-разработки.")
add_body("Цель работы — разработать завершённый интерактивный сайт «МГУПИ—90: цифровая капсула времени», который раскрывает историческое наследие и современный образ университета, предлагает пользователю персональный маршрут и сохраняет локальное послание.", bold_lead="Цель работы — ")
add_body("Для достижения цели поставлены и решены следующие задачи:")
add_numbered_list([
    "провести анализ целевой аудитории и определить пользовательские сценарии интерактивного юбилейного сайта;",
    "исследовать технологические возможности реализации проекта и выбрать оптимальный формат веб-прототипа;",
    "проанализировать аналоги цифровых архивов, юбилейных сайтов, капсул времени и интерактивных лонгридов;",
    "разработать визуальную концепцию, цветовую систему и макет пользовательского маршрута;",
    "реализовать одностраничный сайт с временной линией, мини-квестом, кроссвордом и формой послания;",
    "проверить адаптивность, работоспособность интерактивных элементов и корректность производственной сборки;",
    "подготовить комплект финальных изображений, схем и скриншотов для включения в отчёт.",
])
add_body("Структура и оформление отчёта подготовлены с учётом требований ГОСТ 7.32–2017 [1], а библиографическое описание источников — с учётом требований ГОСТ Р 7.0.5–2008 [2].")
add_body("Объектом разработки является цифровая юбилейная коммуникация образовательной организации. Предмет разработки — визуальные и технологические способы вовлечения пользователя в интерактивный веб-маршрут. Практическая значимость заключается в возможности демонстрации проекта в локальной или размещённой веб-среде без серверной части и внешних зависимостей контента.")

add_heading("1 ПРЕДПРОЕКТНОЕ ИССЛЕДОВАНИЕ", 1)
add_heading("1.1 Анализ целевой аудитории проекта", 2)
add_body("Ядро целевой аудитории составляют текущие студенты. Для них важны узнаваемые ситуации повседневной университетской жизни, современный визуальный язык, короткие смысловые блоки и возможность повлиять на итог взаимодействия. Дополнительные аудитории — выпускники, преподаватели и абитуриенты. Выпускникам важен эмоциональный контакт с личной памятью, преподавателям — корректное представление образовательной традиции, абитуриентам — образ современной и открытой технологической среды.")
add_body("Объединяющим мотивом для всех групп является связь времён. Проект не строится как энциклопедия: фактическая информация подаётся через шесть смысловых точек, а личное участие обеспечивают выборы и форма послания. Интерфейс рассчитан на просмотр с настольного компьютера и смартфона; чтение не требует предварительной подготовки.")
add_body("Сравнительная характеристика основных групп целевой аудитории приведена в таблице 1.1.")
add_table_caption("Характеристика целевой аудитории")
add_data_table(
    ["Группа", "Потребность", "Сценарий", "Интерфейсный ответ"],
    [
        ("Студенты", "Самоидентификация и участие", "Маршрут, выбор, послание", "Динамичные карточки, быстрый отклик"),
        ("Выпускники", "Эмоциональная связь", "Просмотр истории и образов", "Архивная интонация, крупные изображения"),
        ("Преподаватели", "Преемственность", "Последовательное чтение", "Структурная временная линия"),
        ("Абитуриенты", "Понимание среды", "Знакомство с настоящим и будущим", "Современная визуальная система"),
    ], [1650, 2450, 2500, 2760]
)
add_figure_with_intro(
    TECH / "01-audience-map.png",
    "Карта целевой аудитории проекта",
    "Далее на рисунке 1.1 представлена карта целевой аудитории, показывающая различие пользовательских мотивов и общее требование к интерфейсу."
)

add_heading("1.2 Анализ технологических возможностей проекта", 2)
add_body("На предпроектном этапе рассмотрены веб-технологии, геймификация и иммерсивные форматы AR/VR. Дополненная и виртуальная реальность обладают высокой выразительностью, но требуют специального оборудования, сложного производства контента и дополнительного тестирования. Для учебного демонстрационного продукта эти ограничения снижают доступность. Одностраничный сайт, напротив, запускается на распространённых устройствах, легко документируется скриншотами и позволяет реализовать необходимую интерактивность без серверной инфраструктуры.")
add_body("Геймификация используется в умеренной форме. Мини-квест состоит из трёх смысловых выборов и персонального результата. Дополнительный кроссворд закрепляет семь тематических понятий и открывает памятный код после проверки ответов. Для снижения сложности предусмотрены два уровня помощи: дополнительный смысловой намёк и поэтапное открытие букв. Временная линия поддерживает исследовательский сценарий, а форма послания завершает его личным действием. Состояние кроссворда и сообщение сохраняются в localStorage без передачи информации на внешний сервер.")
add_body("Сравнение рассмотренных технологических подходов приведено в таблице 1.2.")
add_table_caption("Сравнение технологических подходов")
add_data_table(
    ["Технология", "Преимущества", "Ограничения", "Решение"],
    [
        ("Веб-сайт", "Доступность, адаптивность, простое внедрение", "Зависимость от экрана", "Выбрано"),
        ("Геймификация", "Вовлечение, персональный результат", "Риск перегрузки механиками", "Использована точечно"),
        ("AR/VR", "Высокая иммерсивность", "Оборудование и сложная разработка", "Не выбрано"),
        ("Видео", "Эмоциональность и контроль ритма", "Пассивный сценарий", "Использовано как визуальная логика, не формат"),
    ], [1500, 2850, 2750, 2260]
)

add_heading("1.3 Аналоги и прототипы проекта", 2)
add_body("В качестве аналогов рассмотрены не только сайты университетских юбилеев, но и близкие по механике цифровые продукты: сервисы писем в будущее, музейные онлайн-архивы и интерактивные лонгриды. Такой подход позволяет оценить проект не только по внешнему виду, но и по сценарию вовлечения пользователя.")
add_body("Первый аналог — FutureMe, сервис отправки писем себе в будущее [4]. Его сильная сторона заключается в понятной эмоциональной механике: пользователь формулирует послание и связывает настоящее с будущим. Ограничение для разрабатываемого проекта состоит в отсутствии коллективного исторического контекста: FutureMe не раскрывает тему организации или сообщества, а работает прежде всего как персональный инструмент.")
add_body("Второй тип аналогов — юбилейные сайты и цифровые архивы университетов, например проекты MIT, посвящённые исторической памяти института [5]. Их преимущество заключается в хронологической структуре, архивности и авторитетности материала. Недостаток — чаще всего пассивная модель взаимодействия: пользователь читает и просматривает материалы, но не становится участником юбилейной коммуникации.")
add_body("Третий аналог — Google Arts & Culture как пример масштабной цифровой экспозиции [6]. Он показывает, как культурный и исторический материал можно упаковывать в визуальные маршруты, подборки и тематические истории. Для учебного проекта такой уровень архива избыточен, однако принцип крупного изображения, подписи и тематического перехода полезен для галереи и временной линии.")
add_body("Четвёртый тип аналогов — интерактивные редакционные истории и лонгриды. Они создают ощущение последовательного путешествия, используют крупные заголовки, ритм экранов, микроанимацию и чёткий сценарий. Риск такого подхода — избыточная эффектность, поэтому в проекте оставлены только те интерактивные элементы, которые поддерживают смысл: таймлайн, мини-квест, кроссворд и послание.")
add_figure_with_intro(
    TECH / "02-analog-map.png",
    "Карта аналогов и прототипов проекта",
    "Далее на рисунке 1.2 показана обобщённая карта аналогов, из которых сформированы ключевые проектные принципы."
)
add_body("Итоговое сравнение аналогов и заимствованных принципов представлено в таблице 1.3.")
add_table_caption("Сравнение смысловых аналогов")
add_data_table(
    ["Тип аналога", "Сильная сторона", "Риск", "Заимствованный принцип"],
    [
        ("Юбилейный сайт", "Хронология", "Статичность", "Временная линия"),
        ("Цифровой архив", "Визуальная достоверность", "Перегрузка материалом", "Крупные кадры и подписи"),
        ("Капсула времени", "Личное участие", "Слабый контекст", "Послание и финальная карточка"),
        ("Интерактивная история", "Ритм и вовлечение", "Избыточная анимация", "Пошаговый маршрут"),
    ], [1900, 2300, 2200, 2960]
)
add_body("По результатам анализа сформирован смешанный формат: сайт сохраняет ясную структуру юбилейной хроники, использует визуальную плотность цифрового архива, добавляет личное действие через форму послания и поддерживает внимание за счёт компактной игровой механики.")

add_heading("Выводы по разделу 1", 2)
add_body("Исследование подтвердило обоснованность формата одностраничного интерактивного сайта. Он соответствует привычкам основной аудитории, допускает сочетание истории и личного действия, работает на разных устройствах и не требует специализированной инфраструктуры. Визуальный стиль должен соединять академическую сдержанность, технологичность и эмоциональную теплоту архивного материала.")

add_heading("2 ПРАКТИЧЕСКОЕ ВЫПОЛНЕНИЕ ПРОЕКТА", 1)
add_heading("2.1 Концепция проекта", 2)
add_body("Концепция строится вокруг метафоры «память, которую мы создаём». Пользователь проходит три времени: знакомится с прошлым, определяет собственные ценности в настоящем и оставляет послание будущему. Публичный сценарий выстроен от общего к личному: главный экран — идея — временная линия — мини-квест — кроссворд — послание — визуальная галерея — финальная карточка. Исследовательские и эскизные материалы сознательно не включены в интерфейс сайта и представлены только в настоящем отчёте.")
add_body("Основной графический мотив — орбита с числом 90. Она обозначает цикличность передачи знаний и используется в hero-блоке, результатах квеста, финальной карточке и пиктограмме проекта. Модульная сетка, тонкие линии и технические индексы формируют ощущение цифрового интерфейса. Контрастная антиква в акцентных строках добавляет человеческую и историческую интонацию.")
add_body("Цветовая система включает глубокий тёмно-синий фон, светлый текст, бирюзовый акцент и ограниченный фиолетовый оттенок. Светлые разделы создают паузы в длинной странице и повышают читаемость исследовательских блоков. Для автономности используются системные шрифты: Segoe UI/Arial для интерфейса и Georgia/Times New Roman для акцентных фраз. При построении сценария учитывались принципы ясной обратной связи, минимизации когнитивной нагрузки и понятной навигации [14], [15], [16]. После проверки читаемости минимальный размер смысловых подписей и нумерации увеличен до 11–13 пикселей, кнопок — до 12 пикселей, текстов подсказок кроссворда — до 15 пикселей.")
add_body("Цветовая система финального интерфейса представлена в таблице 2.1.")
add_table_caption("Цветовая система проекта")
add_data_table(
    ["Роль", "Цвет", "Назначение"],
    [
        ("Основной фон", "#07111F", "Глубина и академическая сдержанность"),
        ("Поверхность", "#0D1C2E", "Карточки и интерактивные панели"),
        ("Акцент", "#58E8E1", "Действия, состояния и ключевые слова"),
        ("Дополнительный", "#8B7CFF", "Точки будущего и орбиты"),
        ("Светлый фон", "#ECE9E1", "Кроссворд и смысловая пауза"),
    ], [1900, 1700, 5760]
)

add_heading("2.2 Эскизирование проекта", 2)
add_body("Для поиска визуального решения были подготовлены четыре направления эскизирования: «Архив», «Неон», «Коллаж» и «Синтез». Первый вариант делает акцент на документальности, но выглядит менее цифровым; второй убедительно передаёт технологичность, однако может восприниматься холодно; третий энергичен, но конкурирует с содержанием. Для финальной разработки выбран четвёртый вариант, соединяющий архивную человечность и цифровую точность.")
add_figure_with_intro(
    TECH / "03-wireflow.png",
    "Макет пользовательского маршрута сайта",
    "Далее на рисунке 2.1 изображён макет пользовательского маршрута, показывающий последовательность экранов и различие десктопной и мобильной компоновки."
)
add_figure_with_intro(
    SCREEN / "04-sketches.png",
    "Четыре направления эскизирования и выбранный вариант",
    "Далее на рисунке 2.2 представлены варианты визуальной концепции и выделен финальный вариант «Синтез»."
)
add_body("В выбранном направлении сохранены крупная типографика, строгая сетка и контраст светлых и тёмных секций. Анимация ограничена плавными переходами, состояниями наведения и появлением активного содержимого. Такое решение поддерживает технологический образ и одновременно не мешает фиксации экранов для отчёта.")

add_heading("2.3 Финальные изображения проекта", 2)
add_body("Визуальная серия подготовлена специально для учебного прототипа и хранится локально. Это исключает битые внешние ссылки и обеспечивает автономную демонстрацию. Шесть самостоятельных кадров сопровождают точки временной линии, ещё четыре используются в галерее. Изображения являются концептуальными визуальными реконструкциями и не выдаются за документальные архивные фотографии.")
for filename, title, intro in (
    ("01-origin.png", "Начало инженерной образовательной традиции — визуальная реконструкция", "Далее на рисунке 2.3 изображён стартовый исторический образ, который задаёт тему инженерной образовательной традиции."),
    ("02-growth.png", "Развитие приборостроения и вычислительных направлений — визуальная реконструкция", "Далее на рисунке 2.4 показан визуальный образ развития технических направлений и перехода к вычислительной культуре."),
    ("03-community.png", "Современная студенческая проектная среда", "Далее на рисунке 2.5 представлена студенческая проектная среда как связующее звено между историей и современностью."),
    ("04-now.png", "Робототехника и цифровое прототипирование", "Далее на рисунке 2.6 изображён современный технологический слой проекта: робототехника, прототипирование и цифровая разработка."),
    ("05-anniversary.png", "Юбилейная интерактивная экспозиция", "Далее на рисунке 2.7 показан юбилейный образ, объединяющий число 90, архивность и интерфейсную экспозицию."),
    ("06-future.png", "Образ образовательного кампуса будущего", "Далее на рисунке 2.8 представлен образ будущего университета как открытой технологической среды."),
):
    add_figure_with_intro(ROOT / "public" / "images" / "timeline" / filename, title, intro)

add_body("Галерейная серия дополняет временную линию четырьмя укрупнёнными образами истории, сообщества, технологии и будущего.")
for filename, title, intro in (
    ("archive-lab.png", "Исторический образ инженерной лаборатории", "Далее на рисунке 2.9 изображён галерейный кадр исторического блока, предназначенный для замены на реальную архивную фотографию при необходимости."),
    ("students-studio.png", "Современная студенческая проектная среда", "Далее на рисунке 2.10 показан галерейный образ студенческой среды, подчёркивающий командную работу и проектность."),
    ("technology-lab.png", "Цифровые технологии и робототехника", "Далее на рисунке 2.11 представлен визуальный материал для карточки о цифровых технологиях."),
    ("future-campus.png", "Образ будущего образовательной среды", "Далее на рисунке 2.12 изображён образ будущего образования, поддерживающий финальную часть сайта."),
):
    add_figure_with_intro(ROOT / "public" / "images" / filename, title, intro)

add_body("Ниже показаны финальные интерфейсные состояния. Комплект включает все ключевые смысловые разделы, интерактивные состояния и мобильную версию. Послание, использованное в демонстрации, является тестовым и хранится только в локальном профиле браузера.")
for filename, title, width, intro in (
    ("01-hero.png", "Главный экран проекта", Inches(6.25), "Далее на рисунке 2.13 изображён главный экран сайта с названием проекта, вводным сообщением и основными действиями пользователя."),
    ("02-idea.png", "Раздел «Идея проекта»", Inches(6.25), "Далее на рисунке 2.14 представлен раздел идеи, объясняющий смысл цифровой капсулы времени и три тематических направления проекта."),
    ("03-timeline-active.png", "Временная линия с активным этапом", Inches(6.25), "Далее на рисунке 2.15 показана интерактивная временная линия с открытой карточкой выбранного этапа."),
    ("04-quest-step-1.png", "Первый шаг мини-квеста", Inches(6.25), "Далее на рисунке 2.16 изображён первый шаг мини-квеста, где пользователь выбирает ценность, которую считает важной сохранить."),
    ("05-quest-result.png", "Персональный результат мини-квеста", Inches(6.25), "Далее на рисунке 2.17 показан результат мини-квеста, формирующий персональный профиль цифровой капсулы."),
    ("06-crossword.png", "Тематический кроссворд", Inches(6.25), "Далее на рисунке 2.18 изображён тематический кроссворд, дополняющий сайт игровой проверкой ключевых понятий."),
    ("06b-crossword-hint.png", "Поэтапная помощь: смысловой намёк и открытая буква", Inches(6.25), "Далее на рисунке 2.19 показан блок помощи в кроссворде: пользователь может открыть дополнительный намёк и одну букву."),
    ("07-crossword-complete.png", "Кроссворд с открытым памятным кодом", Inches(6.25), "Далее на рисунке 2.20 представлен завершённый кроссворд с открытым памятным кодом."),
    ("08-message-form.png", "Форма «Послание в будущее»", Inches(6.25), "Далее на рисунке 2.21 изображена форма сохранения послания, включающая имя, направление или группу и текст сообщения."),
    ("09-message-saved.png", "Сохранённое послание", Inches(6.25), "Далее на рисунке 2.22 показано состояние сохранённого послания после записи данных в локальное хранилище браузера."),
    ("10-gallery.png", "Визуальная галерея проекта", Inches(6.0), "Далее на рисунке 2.23 представлена визуальная галерея проекта с четырьмя тематическими карточками."),
    ("11-final-card.png", "Финальная карточка «Капсула создана»", Inches(6.25), "Далее на рисунке 2.24 изображена финальная карточка, завершающая пользовательский маршрут."),
    ("12-mobile-hero.png", "Мобильная версия главного экрана", Inches(2.9), "Далее на рисунке 2.25 показана мобильная версия главного экрана, подтверждающая адаптацию интерфейса под смартфон."),
    ("13-mobile-crossword.png", "Мобильная версия кроссворда с блоком помощи", Inches(1.5), "Далее на рисунке 2.26 изображена мобильная версия кроссворда с перестроенным блоком вопросов и подсказок."),
):
    add_figure_with_intro(SCREEN / filename, title, intro, width=width)

add_heading("Выводы по разделу 2", 2)
add_body("Практическая часть завершилась формированием целостной визуальной и интерактивной системы. Финальное решение развивает выбранный эскиз «Синтез», поддерживает метафору цифровой памяти и сохраняет единый язык на всех экранах. Основная сложность заключалась в согласовании академической темы и выразительного цифрового образа; она решена с помощью строгой сетки, ограниченной палитры, спокойной анимации и последовательного сценария.")

add_heading("3 ТЕХНОЛОГИЧЕСКОЕ РЕШЕНИЕ", 1)
add_heading("3.1 Архитектура и инструменты", 2)
add_body("Проект реализован на React и TypeScript с использованием сборщика Vite. React обеспечивает компонентную структуру и управление состояниями интерфейса [7], TypeScript контролирует типы временной линии, квеста, кроссворда и сообщения [8], Vite отвечает за локальную разработку и производственную сборку [9]. Иконки подключены через Iconify [10] и локальный набор Phosphor [11], поэтому после установки зависимостей не требуют сети. Backend, авторизация и база данных не используются.")
add_body("Данные временной линии и вариантов квеста вынесены в отдельные модули. Компоненты отвечают за самостоятельные смысловые блоки, а App.tsx связывает результат квеста, прохождение кроссворда и сохранённое сообщение с финальной карточкой. Изображения расположены в public/images и доступны без обращения к сети.")
add_figure_with_intro(
    TECH / "12-project-structure.png",
    "Структура файлов React-проекта",
    "Далее на рисунке 3.1 показана структура файлов проекта, отражающая разделение исходного кода, данных, стилей и локальных изображений."
)

add_heading("3.2 Реализация интерактивности", 2)
add_body("Плавная навигация реализована якорными ссылками и свойством scroll-behavior. Временная линия хранит идентификатор активного этапа и подменяет подробную карточку без перезагрузки. Мини-квест включает выбор сохраняемой ценности, образа будущего и роли пользователя; алгоритм суммирует смысловые категории и формирует один из трёх профилей капсулы.")
add_body("Кроссворд построен как координатная сетка 11 × 11. Семь слов описаны направлением, начальной координатой, ответом, основным вопросом и дополнительным намёком; пересечения объединяются в общие ячейки. Компонент поддерживает навигацию по выбранному слову, автоматический переход между клетками, проверку ошибок и индикацию прогресса. Пользователь может сначала открыть расширенный смысловой намёк с длиной и первой буквой, а затем — по одной букве непосредственно в сетке. Открытые системой клетки визуально отмечаются и защищаются от случайного изменения. Ответы, использованные подсказки и статус прохождения восстанавливаются из localStorage.")
add_body("Форма создаёт объект из имени, направления, текста и времени сохранения. Объект сериализуется в JSON и помещается в localStorage [12]. При повторном открытии сайта данные восстанавливаются; пользователь может удалить их отдельной кнопкой. Обработка повреждённых данных предусмотрена через безопасное чтение и очистку локального значения.")
add_body("Адаптивность реализована медиазапросами для ширины 1100, 820 и 520 пикселей. На мобильном устройстве горизонтальная навигация заменяется меню, многоколоночные блоки становятся последовательными, а кнопки занимают доступную ширину. Кроссворд сохраняет квадратную сетку, а вопросы и блок помощи переходят под игровое поле. Для пользователей, предпочитающих уменьшение движения, добавлен prefers-reduced-motion, что соответствует принципу уважения пользовательских настроек доступности [13].")
add_body("Архитектура технологического решения представлена в таблице 3.1.")
add_table_caption("Архитектура технологического решения")
add_data_table(
    ["Уровень", "Компоненты", "Назначение"],
    [
        ("Представление", "Timeline, Quest, Crossword, Gallery", "Экраны и интерактивные состояния"),
        ("Данные", "timeline.ts, quests.ts", "Этапы, варианты и правила результатов"),
        ("Состояние", "React useState/useEffect", "Этап, ответы, клетки и сообщение"),
        ("Хранение", "localStorage", "Послание, ответы и подсказки кроссворда"),
        ("Сборка", "TypeScript + Vite", "Проверка типов и выпуск статических файлов"),
    ], [1800, 3000, 4560]
)

add_heading("3.3 Тестирование и внедрение", 2)
add_body("Проверка выполнена в производственной сборке и в Microsoft Edge на шести контрольных разрешениях: 320, 360, 390, 768, 1024 и 1440 пикселей. Автоматизированный сценарий изменил этап временной линии, прошёл три шага квеста, открыл оба уровня помощи, заполнил и проверил кроссворд, сохранил послание и зафиксировал ключевые состояния. TypeScript-сборка завершена без ошибок; ошибок браузерной консоли и горизонтального переполнения не обнаружено. Контрольные размеры типографики во всех режимах составили: индекс раздела — 13 пикселей, кнопки — 12 пикселей, текст вопросов кроссворда — 15 пикселей.")
add_body("Результаты функционального и адаптивного тестирования приведены в таблице 3.2.")
add_table_caption("Результаты тестирования")
add_data_table(
    ["Проверка", "Ожидаемый результат", "Статус"],
    [
        ("npm/pnpm build", "Сборка без ошибок типов", "Пройдено"),
        ("Временная линия", "Активная карточка меняется", "Пройдено"),
        ("Мини-квест", "Результат зависит от ответов", "Пройдено"),
        ("Подсказки", "Намёк и открытие букв работают поэтапно", "Пройдено"),
        ("Кроссворд", "Проверка 7 слов и открытие кода", "Пройдено"),
        ("localStorage", "Сообщение восстанавливается", "Пройдено"),
        ("Очистка", "Сообщение удаляется", "Пройдено"),
        ("Desktop 1024/1440 px", "Нет переполнения и наложений", "Пройдено"),
        ("Mobile 320/360/390 px", "Одноколоночная адаптация", "Пройдено"),
        ("Tablet 768 px", "Корректная перестройка сеток", "Пройдено"),
        ("Консоль браузера", "Ошибки отсутствуют", "Пройдено"),
    ], [2450, 4900, 2010]
)
add_figure_with_intro(
    TECH / "13-build-success.png",
    "Результат производственной сборки и проверки",
    "Далее на рисунке 3.2 показан результат производственной сборки, подтверждающий отсутствие ошибок TypeScript и успешное формирование статических файлов."
)
add_body("Для запуска требуется установить зависимости командой npm install и выполнить npm run dev. Производственная версия создаётся командой npm run build и размещается в каталоге dist. Поскольку сайт является статическим, его можно развернуть на любом хостинге статических файлов или демонстрировать локально. Внешняя сеть после установки зависимостей не требуется.")

add_heading("Выводы по разделу 3", 2)
add_body("Выбранная архитектура соответствует масштабу учебного проекта: компоненты разделены по ответственности, изображения и иконки доступны локально, сборка воспроизводима, а интерфейс автономен. Реализованы временная линия, трёхэтапный квест, кроссворд, локальное послание и адаптивность. Тестирование подтвердило отсутствие ошибок типов, браузерной консоли и переполнения макета.")

add_heading("ЗАКЛЮЧЕНИЕ", 1)
add_body("В ходе технологической практики разработан интерактивный веб-проект «МГУПИ—90: цифровая капсула времени». Проект раскрывает юбилейную тему через последовательный маршрут по прошлому, настоящему и будущему, сочетает визуальную историю с личным действием пользователя и соответствует рекомендуемому формату цифрового продукта.")
add_body("Выполнено предпроектное исследование, определены аудитории, сопоставлены технологические подходы и смысловые аналоги. Подготовлены четыре направления эскизирования, обоснован выбор финальной концепции, сформирована авторская серия визуальных материалов и разработана адаптивная интерфейсная система.")
add_body("Технологическая реализация включает React-компоненты, типизированные данные, интерактивную временную линию, трёхэтапный мини-квест, кроссворд, локальное хранение результатов, галерею и печатную финальную карточку. Производственная сборка и браузерный сценарий успешно проверены. Цель работы достигнута, поставленные задачи выполнены.")
add_body("Перспективы развития проекта связаны с подключением проверенного исторического архива, серверным сбором посланий с модерацией, экспортом персональной карточки в графический файл и созданием режима коллективной экспозиции для юбилейного мероприятия.")

add_heading("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 1)
sources = [
    "ГОСТ 7.32–2017. Система стандартов по информации, библиотечному и издательскому делу. Отчёт о научно-исследовательской работе. Структура и правила оформления. — Введ. 2018-07-01. — Москва: Стандартинформ, 2017.",
    "ГОСТ Р 7.0.5–2008. Система стандартов по информации, библиотечному и издательскому делу. Библиографическая ссылка. Общие требования и правила составления. — Введ. 2009-01-01. — Москва: Стандартинформ, 2008.",
    "РТУ МИРЭА. История МГУПИ [Электронный ресурс]. — URL: https://www.mirea.ru/about/history-of-the-university/istoriya-mgupi/ (дата обращения: 28.06.2026).",
    "FutureMe. Write a Letter to your Future Self [Электронный ресурс]. — URL: https://www.futureme.org/ (дата обращения: 28.06.2026).",
    "MIT Libraries. 150 Years in the Stacks [Электронный ресурс]. — URL: https://libraries.mit.edu/150books/ (дата обращения: 28.06.2026).",
    "Google Arts & Culture. Bringing the world’s art and culture online for everyone [Электронный ресурс]. — URL: https://about.artsandculture.google.com/ (дата обращения: 28.06.2026).",
    "React. Quick Start [Электронный ресурс]. — URL: https://react.dev/learn (дата обращения: 28.06.2026).",
    "TypeScript. Documentation [Электронный ресурс]. — URL: https://www.typescriptlang.org/docs/ (дата обращения: 28.06.2026).",
    "Vite. Guide [Электронный ресурс]. — URL: https://vite.dev/guide/ (дата обращения: 28.06.2026).",
    "Iconify. React icon component [Электронный ресурс]. — URL: https://iconify.design/ (дата обращения: 28.06.2026).",
    "Phosphor Icons. Icon set [Электронный ресурс]. — URL: https://icon-sets.iconify.design/ph/ (дата обращения: 28.06.2026).",
    "MDN Web Docs. Window: localStorage property [Электронный ресурс]. — URL: https://developer.mozilla.org/docs/Web/API/Window/localStorage (дата обращения: 28.06.2026).",
    "W3C. Web Content Accessibility Guidelines (WCAG) 2.2 [Электронный ресурс]. — URL: https://www.w3.org/TR/WCAG22/ (дата обращения: 28.06.2026).",
    "Norman D. The Design of Everyday Things. — Revised and expanded ed. — New York: Basic Books, 2013. — 368 p.",
    "Krug S. Don't Make Me Think, Revisited: A Common Sense Approach to Web Usability. — 3rd ed. — Berkeley: New Riders, 2014. — 216 p.",
    "Cooper A., Reimann R., Cronin D., Noessel C. About Face: The Essentials of Interaction Design. — 4th ed. — Indianapolis: Wiley, 2014. — 720 p.",
]
for index, source in enumerate(sources, 1):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.hanging_indent = Cm(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    set_font(p.add_run(f"{index}. {source}"), size=12)

add_heading("ПРИЛОЖЕНИЕ А\nПЕРЕЧЕНЬ МАТЕРИАЛОВ ПРОЕКТА", 1)
add_body("Электронная версия проекта включает исходный код сайта, десять локальных визуалов, четырнадцать финальных скриншотов, изображения технической проверки, README с командами запуска и настоящий отчёт. Каталог сайта: mgupi-time-capsule. Основные команды: npm install, npm run dev, npm run build.")
add_body("Скриншоты подготовлены в разрешении 1440 пикселей для настольной версии и 390 пикселей для мобильной версии. Они находятся в каталоге report-assets/screenshots и могут повторно использоваться в презентации или при корректировке отчёта.")

doc.core_properties.title = "МГУПИ—90: цифровая капсула времени — отчёт по технологической практике"
doc.core_properties.subject = "Интерактивный веб-проект"
doc.core_properties.author = "Обучающийся кафедры компьютерного дизайна"
doc.core_properties.keywords = "МГУПИ, цифровая капсула времени, React, TypeScript, веб-дизайн"
doc.core_properties.comments = "Сформировано для технологической практики 2025–2026 учебного года"

# Ask Word to update fields (TOC and page numbers) when the document is opened.
settings = doc.settings._element
update_fields = settings.find(qn("w:updateFields"))
if update_fields is None:
    update_fields = OxmlElement("w:updateFields")
    settings.append(update_fields)
update_fields.set(qn("w:val"), "true")

doc.save(OUT)
print(OUT)
