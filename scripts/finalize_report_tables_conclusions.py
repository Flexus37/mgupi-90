from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


CONCLUSION_TEXTS = {
    "1": [
        (
            "В результате предпроектного исследования была определена содержательная и "
            "проектная основа интерактивного сайта «МГУПИ-90: цифровая капсула времени». "
            "Анализ целевой аудитории показал, что проект должен быть понятен нескольким "
            "группам пользователей: студентам, выпускникам, абитуриентам и представителям "
            "профессиональной среды. Для студентов важны интерактивность, быстрый вход в "
            "сценарий и возможность оставить личный след; для выпускников — эмоциональная "
            "связь с историей университета; для абитуриентов — ясный и современный образ "
            "образовательной среды. Поэтому в качестве базового принципа была выбрана не "
            "статичная информационная страница, а маршрут, в котором пользователь постепенно "
            "переходит от истории к современности и будущему."
        ),
        (
            "Рассмотрение технологических возможностей позволило уточнить границы учебного "
            "прототипа. Для решения задачи достаточно клиентской веб-разработки без серверной "
            "части: интерфейс реализуется средствами React, TypeScript и Vite, а пользовательское "
            "послание сохраняется локально в браузере. Такой подход соответствует демонстрационному "
            "характеру проекта: сайт можно запустить на учебном компьютере, показать преподавателю "
            "без подключения к базе данных и использовать как источник экранных материалов для "
            "отчета. Одновременно выбранный стек дает возможность реализовать интерактивную "
            "временную линию, мини-игру, кроссворд и форму послания без усложнения архитектуры."
        ),
        (
            "Анализ аналогов и прототипов показал, что наиболее убедительные цифровые проекты "
            "юбилейной тематики соединяют три качества: архивность, персональное участие и "
            "визуальную навигацию. У проектов формата FutureMe была заимствована идея личного "
            "послания в будущее; у музейных и образовательных цифровых экспозиций — принцип "
            "структурированной подачи исторических материалов; у интерактивных лонгридов — "
            "сценарность, последовательность экранов и мягкое вовлечение пользователя. На основе "
            "этого анализа сформирован собственный проектный принцип: сайт не просто сообщает "
            "факты о юбилее, а предлагает пользователю пройти небольшой символический путь."
        ),
        (
            "Таким образом, первый раздел подтвердил актуальность выбранной концепции и задал "
            "критерии для дальнейшей разработки. Будущий интерфейс должен быть академичным, но "
            "не перегруженным; визуально современным, но связанным с темой исторической памяти; "
            "интерактивным, но доступным для быстрого понимания. Эти выводы стали основанием для "
            "перехода к практическому этапу: разработке композиции сайта, пользовательского "
            "маршрута, игровых элементов и визуальных материалов, которые раскрывают идею "
            "цифровой капсулы времени."
        ),
    ],
    "2": [
        (
            "В ходе практического выполнения проекта была сформирована целостная концепция "
            "одностраничного интерактивного сайта, в котором каждый раздел выполняет отдельную "
            "смысловую функцию. Главный экран вводит пользователя в тему юбилея и задает образ "
            "цифровой капсулы времени; блок идеи объясняет назначение проекта; временная линия "
            "показывает развитие образовательной традиции; игровые элементы вовлекают пользователя "
            "в активное взаимодействие; форма послания связывает индивидуальный опыт с образом "
            "будущего. Благодаря такой последовательности сайт воспринимается не набором отдельных "
            "блоков, а единым маршрутом."
        ),
        (
            "Особое внимание на практическом этапе было уделено пользовательскому сценарию. "
            "Маршрут построен так, чтобы пользователь сначала получил общее эмоциональное "
            "впечатление, затем познакомился с историко-смысловыми точками, после этого прошел "
            "небольшую игровую проверку и в завершение оставил собственное послание. Мини-квест "
            "и кроссворд усиливают вовлечение, потому что переводят тему юбилея из пассивного "
            "чтения в действие. При этом игровые элементы не противоречат академическому характеру "
            "проекта: они работают как способ закрепления ключевых понятий — история, технологии, "
            "студенческое сообщество, будущее образования."
        ),
        (
            "Визуальное решение разрабатывалось с учетом выбранного образа: темная основа, "
            "светлая типографика, бирюзово-синие акценты, карточная структура и технологичные "
            "графические элементы. Такая стилистика позволяет объединить тему университетской "
            "истории с ощущением современной цифровой среды. При проектировании были подготовлены "
            "схемы целевой аудитории, карта аналогов, маршрут пользователя и макеты интерфейсных "
            "экранов. Эти материалы помогают показать не только итоговый сайт, но и логику "
            "проектных решений: от исследования аудитории к структуре и визуальному языку продукта."
        ),
        (
            "Итогом второго раздела стало практическое обоснование интерфейса и его основных "
            "разделов. Разработанный сайт решает сразу несколько задач: представляет юбилейную "
            "тему в современном формате, демонстрирует интерактивные возможности веб-среды, "
            "создает набор экранов для отчетной фиксации результата и оставляет место для "
            "дальнейшего наполнения реальными фотографиями. Таким образом, практический этап "
            "подтвердил жизнеспособность выбранной концепции и подготовил проект к технической "
            "реализации."
        ),
    ],
    "3": [
        (
            "Технологическое решение проекта основано на использовании React, TypeScript и Vite, "
            "что позволило создать современный, модульный и удобный для дальнейшей поддержки "
            "учебный веб-прототип. Компонентная структура делает интерфейс прозрачным: отдельные "
            "части сайта вынесены в самостоятельные блоки, данные временной линии и интерактивных "
            "заданий хранятся отдельно от разметки, а стили собраны в общей системе оформления. "
            "Такой подход облегчает проверку проекта, замену изображений, расширение текстов и "
            "добавление новых интерактивных элементов без переработки всего сайта."
        ),
        (
            "В рамках реализации были выполнены ключевые функциональные требования: плавная "
            "навигация по разделам, активное состояние элементов временной линии, мини-квест с "
            "расчетом результата, кроссворд с подсказками, форма послания в будущее, сохранение "
            "данных в localStorage и отображение сохраненного сообщения после перезагрузки страницы. "
            "Отсутствие серверной части является осознанным решением: для учебного демонстрационного "
            "проекта важнее стабильность локального запуска, простота проверки и автономность, чем "
            "полноценная система хранения пользовательских данных."
        ),
        (
            "Отдельным результатом технического этапа стала адаптивная верстка. Интерфейс рассчитан "
            "на просмотр как на настольном экране, так и на мобильном устройстве: сетки перестраиваются, "
            "карточки становятся более компактными, навигация и игровые элементы сохраняют читаемость. "
            "Это особенно важно для проекта, ориентированного на студентов и абитуриентов, которые "
            "часто взаимодействуют с образовательными материалами со смартфона. Кроме того, использование "
            "локальных данных и внутренних визуальных заглушек позволяет сайту корректно работать без "
            "постоянной зависимости от внешних ресурсов."
        ),
        (
            "Таким образом, третий раздел подтверждает, что выбранное технологическое решение "
            "соответствует задачам проекта и обеспечивает достаточный уровень качества для учебной "
            "практики. Сайт можно запускать локально, собирать в итоговую версию, демонстрировать "
            "как интерактивный прототип и использовать для подготовки иллюстративных материалов "
            "отчета. Техническая реализация поддерживает основную идею работы: цифровая капсула "
            "времени представлена не только как визуальный образ, но и как рабочий интерактивный "
            "продукт, объединяющий содержание, дизайн и пользовательское действие."
        ),
    ],
}


def find_source() -> Path:
    base = Path.home() / "Downloads" / "Telegram Desktop"
    candidates = [
        p
        for p in base.glob("*.docx")
        if "ready" in p.name.lower()
        and "90" in p.name
        and not p.name.startswith("~$")
        and "_tables_conclusions" not in p.stem
    ]
    if not candidates:
        raise FileNotFoundError(f"No matching ready DOCX found in {base}")
    return max(candidates, key=lambda p: p.stat().st_mtime)


def set_cell_shading(cell, fill: str = "FFFFFF") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for shd in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(shd)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_cell_borders(cell, color: str = "000000", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, margin_dxa: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "left", "bottom", "right"):
        element = tc_mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(margin_dxa))
        element.set(qn("w:type"), "dxa")


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    existing = tr_pr.find(qn("w:tblHeader"))
    if existing is None:
        hdr = OxmlElement("w:tblHeader")
        hdr.set(qn("w:val"), "true")
        tr_pr.append(hdr)


def style_tables(document: Document) -> int:
    table_count = 0
    for table in document.tables:
        table_count += 1
        try:
            table.style = "Table Grid"
        except Exception:
            pass
        table.alignment = WD_ALIGN_PARAGRAPH.LEFT
        table.allow_autofit = True
        if table.rows:
            set_repeat_header(table.rows[0])

        for r_idx, row in enumerate(table.rows):
            for cell in row.cells:
                set_cell_shading(cell, "FFFFFF")
                set_cell_borders(cell, "000000", "6")
                set_cell_margins(cell, 120)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(3)
                    paragraph.paragraph_format.line_spacing = 1.15
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        run.font.name = "Times New Roman"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                        if r_idx == 0:
                            run.bold = True
    return table_count


def iter_body_elements(document: Document):
    body = document._body._element
    return list(body)


def paragraph_text(element) -> str:
    if element.tag != qn("w:p"):
        return ""
    return "".join(t.text or "" for t in element.iter(qn("w:t"))).strip()


def paragraph_style_id(element) -> str:
    if element.tag != qn("w:p"):
        return ""
    p_pr = element.find(qn("w:pPr"))
    if p_pr is None:
        return ""
    p_style = p_pr.find(qn("w:pStyle"))
    if p_style is None:
        return ""
    return p_style.get(qn("w:val"), "")


def is_section_heading(element) -> bool:
    text = paragraph_text(element)
    style_id = paragraph_style_id(element).lower()
    if style_id.startswith("heading") or style_id.startswith("заголовок"):
        return True
    if text in {"ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ", "СПИСОК ИСПОЛЬЗУЕМЫХ ИСТОЧНИКОВ"}:
        return True
    if text.startswith(("1 ", "2 ", "3 ", "1.", "2.", "3.")):
        return True
    if text.startswith("Выводы по разделу"):
        return True
    return False


def clear_elements_between(body, start_index: int, end_index: int) -> None:
    elements = list(body)
    for element in elements[start_index + 1 : end_index]:
        body.remove(element)


def make_body_paragraph(text: str):
    p = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")

    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "120")
    spacing.set(qn("w:line"), "360")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)

    ind = OxmlElement("w:ind")
    ind.set(qn("w:firstLine"), "708")
    p_pr.append(ind)

    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "both")
    p_pr.append(jc)

    p.append(p_pr)

    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:eastAsia"), "Times New Roman")
    fonts.set(qn("w:cs"), "Times New Roman")
    r_pr.append(fonts)

    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "28")
    r_pr.append(size)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    r_pr.append(color)

    r.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def replace_conclusions(document: Document) -> dict[str, int]:
    body = document._body._element
    replaced: dict[str, int] = {}

    for section_number, paragraphs in CONCLUSION_TEXTS.items():
        elements = list(body)
        target_index = None
        for idx, element in enumerate(elements):
            text = paragraph_text(element)
            normalized = " ".join(text.split())
            if normalized.startswith(f"Выводы по разделу {section_number}"):
                target_index = idx
                break
        if target_index is None:
            continue

        end_index = len(elements)
        for idx in range(target_index + 1, len(elements)):
            element = elements[idx]
            if element.tag == qn("w:sectPr"):
                end_index = idx
                break
            if is_section_heading(element):
                end_index = idx
                break

        clear_elements_between(body, target_index, end_index)
        heading_element = list(body)[target_index]
        anchor = heading_element
        for text in paragraphs:
            new_p = make_body_paragraph(text)
            anchor.addnext(new_p)
            anchor = new_p
        replaced[section_number] = len(paragraphs)

    return replaced


def conclusion_word_counts(document: Document) -> dict[str, int]:
    counts: dict[str, int] = {}
    elements = iter_body_elements(document)
    for section_number in CONCLUSION_TEXTS:
        start = None
        for idx, element in enumerate(elements):
            if paragraph_text(element).startswith(f"Выводы по разделу {section_number}"):
                start = idx
                break
        if start is None:
            counts[section_number] = 0
            continue
        words = 0
        for element in elements[start + 1 :]:
            if is_section_heading(element):
                break
            words += len(paragraph_text(element).split())
        counts[section_number] = words
    return counts


def main() -> None:
    src = find_source()
    out = src.with_name(f"{src.stem}_tables_conclusions.docx")

    document = Document(src)
    table_count = style_tables(document)
    replaced = replace_conclusions(document)
    document.save(out)

    check_doc = Document(out)
    counts = conclusion_word_counts(check_doc)

    print(f"SOURCE={src}")
    print(f"OUTPUT={out}")
    print(f"TABLES_STYLED={table_count}")
    print(f"CONCLUSIONS_REPLACED={replaced}")
    print(f"CONCLUSION_WORD_COUNTS={counts}")


if __name__ == "__main__":
    main()
